from docx import Document
from docx.shared import Inches
from io import BytesIO
from datetime import datetime

def generate_transfer_document(lot_title: str, seller_name: str, buyer_name: str, final_price: float, document_type: str) -> BytesIO:
    """
    Генерирует документ подтверждения передачи прав на лот в формате .docx.
    """
    document = Document()
    document.add_heading('Документ подтверждения передачи прав на лот', level=1)

    document.add_paragraph(f"Дата составления: {datetime.now().strftime('%d.%m.%Y')}")
    document.add_paragraph(f"Тип документа: {document_type.capitalize()}")

    document.add_heading('Информация о лоте', level=2)
    document.add_paragraph(f"Название лота: {lot_title}")
    document.add_paragraph(f"Финальная стоимость: {final_price} ₽")

    document.add_heading('Стороны сделки', level=2)
    document.add_paragraph(f"Продавец: {seller_name}")
    document.add_paragraph(f"Покупатель: {buyer_name}")

    document.add_heading('Условия передачи', level=3)
    if document_type == 'jewelry':
        document.add_paragraph("Настоящим подтверждается передача ювелирного изделия в соответствии с условиями аукциона. Изделие проверено и соответствует описанию.")
    elif document_type == 'historical':
        document.add_paragraph("Настоящим подтверждается передача исторически ценного предмета. Ответственность за сохранность и подлинность переходит к покупателю.")
    else: # standard
        document.add_paragraph("Настоящим подтверждается передача лота в соответствии с условиями аукциона. Претензии по качеству и состоянию принимаются в течение 3 дней с момента передачи.")

    document.add_paragraph("\nПодписи сторон:")
    document.add_paragraph("Продавец: _________________________")
    document.add_paragraph("Покупатель: _________________________")

    # сохр. документ в байтовый поток
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
